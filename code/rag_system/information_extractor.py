import os
import json
from pathlib import Path
import logging
import PyPDF2
import docx
import re
import pdfplumber

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("rag_system.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('information_extractor')

class InformationExtractor:
    def __init__(self, crawl_directory):
        """Initialize the extractor with basic tracking and file source mapping."""
        self.crawl_directory = Path(crawl_directory)
        self.extracted_data = []
        self.successful_files = []
        self.failed_files = []
        
        self.extraction_stats = {
            'success': 0,
            'failed': 0,
            'by_filetype': {
                'txt': 0,
                'pdf': 0,
                'docx': 0
            }
        }
        
        # Load file source mapping if available
        self.file_source_mapping = {}
        source_mapping_files = list(Path(crawl_directory).glob("level_*/file_sources.json"))
        for mapping_file in source_mapping_files:
            try:
                with open(mapping_file, 'r') as f:
                    level_mapping = json.load(f)
                    self.file_source_mapping.update(level_mapping)
            except Exception as e:
                logger.warning(f"Error loading source mapping {mapping_file}: {str(e)}")
        
        logger.info(f"Loaded source information for {len(self.file_source_mapping)} files")
        logger.info(f"Initialized extractor with directory: {crawl_directory}")
    
    def extract_all(self):
        """Extract content from txt, pdf, and docx files."""
        level_dirs = [d for d in self.crawl_directory.iterdir() 
                     if d.is_dir() and d.name.startswith('level_')]
        
        for level_dir in level_dirs:
            logger.info(f"Processing level: {level_dir.name}")
            # Extract web page content
            self._extract_web_content(level_dir)
            # Extract PDF files
            self._extract_pdf_files(level_dir)
            # Extract DOCX files
            self._extract_doc_files(level_dir)
        
        # Generate extraction report
        report_path = Path(self.crawl_directory).parent / "extraction_report.json"
        self.save_extraction_report(report_path)
        
        logger.info(f"Completed extraction. Total documents: {len(self.extracted_data)}")
        return self.extracted_data
    
    def _extract_web_content(self, level_dir):
        """Extract content from scraped web pages with simplified structure."""
        text_content_dir = level_dir / "text_content"
        if not text_content_dir.exists():
            logger.warning(f"No text_content directory found in {level_dir}")
            return

        text_files = [f for f in text_content_dir.iterdir() 
                     if f.is_file() and f.suffix == '.txt']

        for text_file in text_files:
            try:
                with open(text_file, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()

                # Extract URL from header
                url = ""
                for line in content.split('\n')[:10]:
                    if line.startswith("URL:"):
                        url = line.replace("URL:", "").strip()
                        break

                # Split content at the separator line
                parts = content.split("---")
                main_content = parts[-1] if len(parts) > 1 else content

                # Clean up the content - preserve paragraph structure
                cleaned_lines = []
                current_paragraph = []

                for line in main_content.split('\n'):
                    line = line.strip()
                    if line and not line.startswith("==="):
                        current_paragraph.append(line)
                    elif current_paragraph:
                        cleaned_lines.append(" ".join(current_paragraph))
                        current_paragraph = []

                # Add the last paragraph if exists
                if current_paragraph:
                    cleaned_lines.append(" ".join(current_paragraph))

                # Format as a clean document with paragraphs
                processed_content = "\n\n".join(cleaned_lines)

                self.extracted_data.append({
                    'content': processed_content,
                    'metadata': {
                        'source': url,
                        'source_url': url,
                        'type': 'web_page',
                        'filename': text_file.name,
                        'level': level_dir.name
                    }
                })

                self.successful_files.append(text_file.name)
                self.extraction_stats['success'] += 1
                self.extraction_stats['by_filetype']['txt'] = self.extraction_stats['by_filetype'].get('txt', 0) + 1
                logger.info(f"Extracted web content from: {text_file.name}")

            except Exception as e:
                self.failed_files.append({
                    'filename': text_file.name,
                    'error': str(e),
                    'file_type': 'txt'
                })
                logger.error(f"Error extracting from {text_file}: {str(e)}")
                self.extraction_stats['failed'] += 1
    
    def _extract_pdf_files(self, level_dir):
        """Extract text from PDF files with enhanced page tracking and financial information detection."""
        files_dir = level_dir / "downloaded_files"
        if not files_dir.exists():
            return
                
        # Find all PDF files, including those with possibly truncated extensions
        pdf_files = list(files_dir.glob('**/*.pdf'))
        potential_truncated_pdfs = list(files_dir.glob('**/*.pd'))
        
        # Fix truncated extensions if found
        for truncated_file in potential_truncated_pdfs:
            fixed_path = Path(str(truncated_file) + 'f')
            if not fixed_path.exists():  # Don't overwrite existing files
                try:
                    truncated_file.rename(fixed_path)
                    logger.info(f"Fixed truncated extension: {truncated_file} → {fixed_path}")
                    pdf_files.append(fixed_path)
                except Exception as e:
                    logger.error(f"Error fixing truncated PDF extension for {truncated_file}: {str(e)}")
        
        # Import pdfplumber if available for improved extraction
        try:
            import pdfplumber
            pdfplumber_available = True
        except ImportError:
            logger.warning("pdfplumber not available. Using PyPDF2 only. Install with: pip install pdfplumber")
            pdfplumber_available = False
        
        # Define financial/budget keywords for special handling
        financial_keywords = [
            "budget", "budgetary", "allocation", "allocate", "allocating", "allocated", 
            "revenue", "expenditure", "spend", "spending", "spent", "financial", "finance",
            "million", "billion", "savings", "saving", "efficiency", "efficiencies", 
            "target", "targets", "cost", "costs", "funding", "fiscal", "investment",
            "£", "$", "€", "percent", "%", "net", "gross", "total", "deficit", "surplus"
        ]
        
        for pdf_file in pdf_files:
            try:
                text = ""
                extraction_method_used = "None"
                content_by_page = {}  # Track content by page number for all pages
                financial_pages = []  # Track which pages have financial info
                
                # Try PyPDF2 first
                try:
                    with open(pdf_file, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        total_pages = len(pdf_reader.pages)
                        
                        # Extract text with page numbers
                        for page_num in range(total_pages):
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            
                            if page_text:
                                # Check for financial/budget terms on this page
                                page_has_financial_info = any(keyword in page_text.lower() for keyword in financial_keywords)
                                
                                # Add page number markers, with special marking for financial pages
                                if page_has_financial_info:
                                    text += f"[Page {page_num + 1} - FINANCIAL/BUDGET INFORMATION]\n{page_text}\n\n"
                                    financial_pages.append(page_num + 1)
                                else:
                                    text += f"[Page {page_num + 1}]\n{page_text}\n\n"
                                
                                # Track content preview for all pages
                                content_by_page[page_num + 1] = page_text[:100] + "..."
                    
                    if len(text.strip()) > 100:
                        extraction_method_used = "PyPDF2"
                    else:
                        logger.warning(f"PyPDF2 extracted minimal text from {pdf_file.name}, trying alternative method")
                        
                except Exception as pdf_error:
                    logger.warning(f"PyPDF2 error with {pdf_file.name}: {str(pdf_error)}")
                
                # Try pdfplumber if PyPDF2 didn't get good results
                if pdfplumber_available and (extraction_method_used == "None" or len(text.strip()) < 500):
                    try:
                        text_plumber = ""
                        content_by_page_plumber = {}
                        financial_pages_plumber = []
                        
                        with pdfplumber.open(pdf_file) as pdf:
                            total_pages = len(pdf.pages)
                            
                            for page_num, page in enumerate(pdf.pages):
                                page_text = page.extract_text()
                                
                                if page_text:
                                    # Check for financial/budget terms on this page
                                    page_has_financial_info = any(keyword in page_text.lower() for keyword in financial_keywords)
                                    
                                    # Add page number markers, with special marking for financial pages
                                    if page_has_financial_info:
                                        text_plumber += f"[Page {page_num + 1} - FINANCIAL/BUDGET INFORMATION]\n{page_text}\n\n"
                                        financial_pages_plumber.append(page_num + 1)
                                    else:
                                        text_plumber += f"[Page {page_num + 1}]\n{page_text}\n\n"
                                    
                                    # Track content preview for all pages
                                    content_by_page_plumber[page_num + 1] = page_text[:100] + "..."
                        
                        # If pdfplumber got better results, use those instead
                        if len(text_plumber.strip()) > len(text.strip()):
                            text = text_plumber
                            content_by_page = content_by_page_plumber
                            financial_pages = financial_pages_plumber
                            extraction_method_used = "pdfplumber"
                            logger.info(f"Used pdfplumber for better extraction from {pdf_file.name}")
                    
                    except Exception as plumber_error:
                        logger.warning(f"pdfplumber error with {pdf_file.name}: {str(plumber_error)}")
                
                # If we have text, add it to the extracted data
                if text.strip():
                    # Get source information from mapping if available
                    source_url = self.file_source_mapping.get(pdf_file.name, {}).get('source_url', '')
                    source_download_url = self.file_source_mapping.get(pdf_file.name, {}).get('download_url', '')
                    
                    # Create metadata with both general page tracking and financial page tracking
                    metadata = {
                        'source': source_url or pdf_file.name,
                        'source_url': source_url,
                        'download_url': source_download_url,
                        'type': 'pdf',
                        'filename': pdf_file.name,
                        'level': level_dir.name,
                        'page_count': total_pages,
                        'extraction_method': extraction_method_used,
                        'pages_with_content': list(content_by_page.keys())  # All pages with content
                    }
                    
                    # Add financial information metadata if found
                    if financial_pages:
                        metadata['contains_financial_info'] = True
                        metadata['financial_pages'] = financial_pages  # Pages with financial info
                    
                    self.extracted_data.append({
                        'content': text,
                        'metadata': metadata
                    })
                    
                    self.successful_files.append(pdf_file.name)
                    self.extraction_stats['success'] += 1
                    self.extraction_stats['by_filetype']['pdf'] = self.extraction_stats['by_filetype'].get('pdf', 0) + 1
                    
                    # Log with appropriate details
                    if financial_pages:
                        logger.info(f"Extracted PDF with financial content from: {pdf_file.name} ({extraction_method_used}, {total_pages} pages, financial pages: {financial_pages})")
                    else:
                        logger.info(f"Extracted PDF content from: {pdf_file.name} ({extraction_method_used}, {total_pages} pages)")
                else:
                    self.failed_files.append({
                        'filename': pdf_file.name,
                        'error': f"No text extracted using {extraction_method_used} (possibly scanned PDF)",
                        'file_type': 'pdf'
                    })
                    logger.warning(f"No text extracted from PDF: {pdf_file.name}")
                    self.extraction_stats['failed'] += 1
                        
            except Exception as e:
                self.failed_files.append({
                    'filename': str(pdf_file),
                    'error': str(e),
                    'file_type': 'pdf'
                })
                logger.error(f"Error extracting from PDF {pdf_file}: {str(e)}")
                self.extraction_stats['failed'] += 1
    
    def _extract_doc_files(self, level_dir):
        """Extract text from DOCX files with enhanced structure and financial information detection."""
        files_dir = level_dir / "downloaded_files"
        if not files_dir.exists():
            return
                
        # Define financial/budget keywords for special handling
        financial_keywords = [
            "budget", "budgetary", "allocation", "allocate", "allocating", "allocated", 
            "revenue", "expenditure", "spend", "spending", "spent", "financial", "finance",
            "million", "billion", "savings", "saving", "efficiency", "efficiencies", 
            "target", "targets", "cost", "costs", "funding", "fiscal", "investment",
            "£", "$", "€", "percent", "%", "net", "gross", "total", "deficit", "surplus"
        ]
        
        # Handle DOCX files
        docx_files = list(files_dir.glob('**/*.docx'))
        
        for docx_file in docx_files:
            try:
                doc = docx.Document(docx_file)
                
                # Track document structure and financial sections
                section_count = len(doc.sections)
                content_by_section = {}
                financial_sections = []
                content_with_headers = []
                
                # Process paragraphs with headings
                current_heading = "Introduction"
                financial_content_found = False
                
                # Process document by paragraphs
                for i, para in enumerate(doc.paragraphs):
                    if para.text:
                        # Identify headings by style
                        if para.style.name.startswith('Heading'):
                            current_heading = para.text
                            heading_level = int(para.style.name.replace('Heading', '')) if para.style.name.replace('Heading', '').isdigit() else 1
                            content_with_headers.append(f"{'#' * heading_level} {para.text}\n\n")
                        else:
                            # Regular paragraph
                            content_with_headers.append(f"{para.text}\n\n")
                        
                        # Track section content (simplified approach)
                        if current_heading not in content_by_section:
                            content_by_section[current_heading] = []
                        content_by_section[current_heading].append(para.text)
                        
                        # Check for financial content
                        if any(keyword in para.text.lower() for keyword in financial_keywords):
                            financial_content_found = True
                            if current_heading not in financial_sections:
                                financial_sections.append(current_heading)
                
                # Process tables - often contain financial data
                tables_content = []
                tables_with_financial_info = []
                
                for i, table in enumerate(doc.tables):
                    table_text = []
                    table_has_financial_info = False
                    
                    for row_idx, row in enumerate(table.rows):
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            row_text.append(cell_text)
                            
                            # Check for financial content in cell
                            if any(keyword in cell_text.lower() for keyword in financial_keywords):
                                table_has_financial_info = True
                        
                        # Add row to table text
                        if row_text:
                            table_text.append(' | '.join(row_text))
                    
                    # Format table with markdown-style formatting
                    if table_text:
                        table_content = f"\n[Table {i+1}]\n" + '\n'.join(table_text)
                        tables_content.append(table_content)
                        content_with_headers.append(table_content + "\n\n")
                        
                        if table_has_financial_info:
                            tables_with_financial_info.append(i+1)
                            financial_content_found = True
                
                # Combine all content with structure
                text = ''.join(content_with_headers)
                
                if text.strip():
                    # Get source information from mapping if available
                    source_url = self.file_source_mapping.get(docx_file.name, {}).get('source_url', '')
                    source_download_url = self.file_source_mapping.get(docx_file.name, {}).get('download_url', '')
                    
                    # Create metadata with structure information
                    metadata = {
                        'source': source_url or docx_file.name,
                        'source_url': source_url,
                        'download_url': source_download_url,
                        'type': 'docx',
                        'filename': docx_file.name,
                        'level': level_dir.name,
                        'section_count': section_count,
                        'paragraph_count': len(doc.paragraphs),
                        'table_count': len(doc.tables),
                        'sections': list(content_by_section.keys())
                    }
                    
                    # Add financial information metadata if found
                    if financial_content_found:
                        metadata['contains_financial_info'] = True
                        metadata['financial_sections'] = financial_sections
                        
                        if tables_with_financial_info:
                            metadata['tables_with_financial_info'] = tables_with_financial_info
                    
                    self.extracted_data.append({
                        'content': text,
                        'metadata': metadata
                    })
                    
                    self.successful_files.append(docx_file.name)
                    self.extraction_stats['success'] += 1
                    self.extraction_stats['by_filetype']['docx'] = self.extraction_stats['by_filetype'].get('docx', 0) + 1
                    
                    if financial_content_found:
                        logger.info(f"Extracted DOCX with financial content from: {docx_file.name} (sections: {financial_sections})")
                    else:
                        logger.info(f"Extracted DOCX content from: {docx_file.name}")
                else:
                    self.failed_files.append({
                        'filename': docx_file.name,
                        'error': "No text extracted from DOCX",
                        'file_type': 'docx'
                    })
                    logger.warning(f"No text extracted from DOCX: {docx_file.name}")
                    self.extraction_stats['failed'] += 1
                        
            except Exception as e:
                self.failed_files.append({
                    'filename': str(docx_file),
                    'error': str(e),
                    'file_type': 'docx'
                })
                logger.error(f"Error extracting from DOCX {docx_file}: {str(e)}")
                self.extraction_stats['failed'] += 1
    
    def save_extraction_report(self, output_path):
        """Save a detailed extraction report including failures."""
        report = {
            'stats': self.extraction_stats,
            'successful_files_count': len(self.successful_files),
            'failed_files_count': len(self.failed_files),
            'failed_files': self.failed_files
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2)
            
        logger.info(f"Saved extraction report to {output_path}")
        return report